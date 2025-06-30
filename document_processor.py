"""
Document processing utilities for EmotiBot
Handles PDF and DOCX document reading and processing
"""
import PyPDF2
from docx import Document
import io
from typing import List, Optional, Dict
import os
import logging

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.logger = logging.getLogger(__name__)
    
    def read_pdf(self, file_path: str) -> Optional[str]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text or None if failed
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                    text += "\n"  # Add newline between pages
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {e}")
            return None
    
    def read_pdf_from_bytes(self, pdf_bytes: bytes) -> Optional[str]:
        """
        Extract text from PDF bytes
        
        Args:
            pdf_bytes: PDF file as bytes
            
        Returns:
            Extracted text or None if failed
        """
        try:
            text = ""
            pdf_file = io.BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
                text += "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error reading PDF from bytes: {e}")
            return None
    
    def read_docx(self, file_path: str) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text or None if failed
        """
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX {file_path}: {e}")
            return None
    
    def read_docx_from_bytes(self, docx_bytes: bytes) -> Optional[str]:
        """
        Extract text from DOCX bytes
        
        Args:
            docx_bytes: DOCX file as bytes
            
        Returns:
            Extracted text or None if failed
        """
        try:
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX from bytes: {e}")
            return None
    
    def read_txt(self, file_path: str, encoding: str = 'utf-8') -> Optional[str]:
        """
        Read text from TXT file
        
        Args:
            file_path: Path to TXT file
            encoding: File encoding (default: utf-8)
            
        Returns:
            File content or None if failed
        """
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as file:
                        return file.read()
                except:
                    continue
            
            self.logger.error(f"Could not decode text file {file_path}")
            return None
            
        except Exception as e:
            self.logger.error(f"Error reading TXT {file_path}: {e}")
            return None
    
    def read_document(self, file_path: str) -> Optional[str]:
        """
        Read document based on file extension
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text or None if failed
        """
        if not os.path.exists(file_path):
            self.logger.error(f"File not found: {file_path}")
            return None
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.read_pdf(file_path)
        elif file_ext == '.docx':
            return self.read_docx(file_path)
        elif file_ext == '.txt':
            return self.read_txt(file_path)
        else:
            self.logger.error(f"Unsupported file format: {file_ext}")
            return None
    
    def read_document_from_bytes(self, file_bytes: bytes, file_type: str) -> Optional[str]:
        """
        Read document from bytes based on file type
        
        Args:
            file_bytes: Document as bytes
            file_type: File type ('pdf', 'docx', 'txt')
            
        Returns:
            Extracted text or None if failed
        """
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return self.read_pdf_from_bytes(file_bytes)
        elif file_type == 'docx':
            return self.read_docx_from_bytes(file_bytes)
        elif file_type == 'txt':
            try:
                return file_bytes.decode('utf-8')
            except:
                return file_bytes.decode('latin-1')
        else:
            self.logger.error(f"Unsupported file type: {file_type}")
            return None
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """
        Split text into chunks for processing
        
        Args:
            text: Text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at word boundary
            if end < len(text):
                # Find the last space before the chunk_size limit
                space_pos = text.rfind(' ', start, end)
                if space_pos > start:
                    end = space_pos
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start < 0:
                start = 0
        
        return chunks
    
    def extract_metadata(self, file_path: str) -> Dict:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dictionary with metadata
        """
        metadata = {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_size': 0,
            'file_type': '',
            'page_count': 0,
            'word_count': 0,
            'character_count': 0
        }
        
        try:
            # Basic file info
            metadata['file_size'] = os.path.getsize(file_path)
            metadata['file_type'] = os.path.splitext(file_path)[1].lower()
            
            # Extract text to count words/characters
            text = self.read_document(file_path)
            if text:
                metadata['character_count'] = len(text)
                metadata['word_count'] = len(text.split())
            
            # Try to get page count for PDF
            if metadata['file_type'] == '.pdf':
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        metadata['page_count'] = len(pdf_reader.pages)
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting metadata from {file_path}: {e}")
            return metadata
    
    def is_supported_format(self, file_path: str) -> bool:
        """
        Check if file format is supported
        
        Args:
            file_path: Path to file
            
        Returns:
            True if supported, False otherwise
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in self.supported_formats
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported file formats
        
        Returns:
            List of supported file extensions
        """
        return self.supported_formats.copy()
